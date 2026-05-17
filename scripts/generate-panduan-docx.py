"""Generate user guide .docx for Arcade Dental ERP (business flow, non-technical)."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUTPUT = Path(__file__).resolve().parent.parent / "Panduan_Penggunaan_Arcade_Dental_ERP.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x29, 0x42)
    return h


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("PANDUAN PENGGUNAAN\nARCADE DENTAL ERP")
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.color.rgb = RGBColor(0x1E, 0x88, 0xE5)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Sistem Manajemen Klinik Gigi — Panduan Alur Bisnis")
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()
    add_para(
        doc,
        "Dokumen ini menjelaskan cara menggunakan website Arcade Dental ERP untuk kegiatan "
        "sehari-hari di klinik. Panduan ditujukan untuk staf klinik (bukan untuk programmer). "
        "Fokusnya adalah alur kerja bisnis: dari kunjungan pasien hingga pembayaran.",
    )

    # 1
    add_heading(doc, "1. Apa itu Arcade Dental ERP?", 1)
    add_para(
        doc,
        "Arcade Dental ERP adalah aplikasi web untuk mengelola operasional klinik gigi. "
        "Semua data disimpan di browser komputer Anda (mode demo), sehingga perubahan tetap "
        "ada setelah halaman di-refresh, selama Anda tidak menghapus data browser.",
    )
    add_bullet(doc, "Mendaftarkan dan mengelola data pasien")
    add_bullet(doc, "Menjadwalkan janji temu dengan dokter")
    add_bullet(doc, "Mencatat perawatan dan layanan")
    add_bullet(doc, "Mengelola stok bahan klinik")
    add_bullet(doc, "Membuat tagihan dan mencatat pembayaran")
    add_bullet(doc, "Melihat laporan dan grafik perkembangan klinik")

    # 2
    add_heading(doc, "2. Cara Membuka Website", 1)
    add_para(doc, "Jalankan aplikasi di komputer yang sudah terpasang programnya, lalu buka browser (Chrome, Edge, dll.) dan ketik alamat:")
    add_para(doc, "http://localhost:3000", bold=True)
    add_para(
        doc,
        "Jika aplikasi di-hosting di server klinik, gunakan alamat yang diberikan oleh tim IT. "
        "Halaman pertama yang muncul adalah Halaman Beranda (profil klinik).",
    )

    # 3
    add_heading(doc, "3. Halaman Beranda (Sebelum Login)", 1)
    add_para(doc, "Di halaman beranda Anda dapat:")
    add_bullet(doc, "Membaca informasi klinik: layanan, tim dokter, dan kontak")
    add_bullet(doc, "Mengganti bahasa: klik ikon globe (🌐) di kanan atas → pilih Bahasa Indonesia atau English")
    add_bullet(doc, "Mengganti tampilan terang/gelap: ikon matahari/bulan di navbar")
    add_bullet(doc, "Masuk ke sistem ERP: tombol “Akses Dashboard ERP” atau menu “Masuk ERP”")

    # 4
    add_heading(doc, "4. Cara Login (Masuk ke Dashboard)", 1)
    add_para(
        doc,
        "Hanya akun yang sudah terdaftar yang bisa masuk. Jika email atau password salah, "
        "akan muncul pesan peringatan merah di pojok layar.",
    )
    add_table(
        doc,
        ["Peran", "Email", "Password", "Kegunaan utama"],
        [
            ["Super Admin", "superadmin@arcadedental.com", "super123", "Akses penuh + pengaturan sistem"],
            ["Admin Klinik", "admin@arcadedental.com", "admin123", "Operasional klinik (tanpa reset data)"],
            ["Dokter", "dokter@arcadedental.com", "dokter123", "Janji temu, pasien, perawatan"],
        ],
    )
    add_para(doc, "Langkah login:", bold=True)
    add_bullet(doc, "Buka halaman login (dari beranda atau menu Masuk ERP)")
    add_bullet(doc, "Isi Email dan Password sesuai tabel di atas")
    add_bullet(doc, "Klik “Masuk Dashboard”")
    add_bullet(doc, "Anda akan diarahkan ke Dashboard Utama sesuai hak akses peran")

    add_para(
        doc,
        "Catatan: Bahasa yang dipilih di halaman beranda tetap dipakai setelah login "
        "(disimpan otomatis). Di halaman login tidak ada pengaturan bahasa terpisah.",
    )

    # 5
    add_heading(doc, "5. Perbedaan Hak Akses per Peran", 1)
    add_table(
        doc,
        ["Menu / Fitur", "Super Admin", "Admin", "Dokter"],
        [
            ["Dashboard (ringkasan)", "✓", "✓", "✓ (tanpa data uang & stok)"],
            ["Janji Temu", "✓", "✓", "✓"],
            ["Pasien", "✓", "✓", "✓"],
            ["Perawatan (master layanan)", "✓", "✓", "✓"],
            ["Dokter & Staff", "✓", "✓", "—"],
            ["Inventori", "✓", "✓", "—"],
            ["Billing / Invoice", "✓", "✓", "—"],
            ["Laporan", "✓", "✓", "—"],
            ["Pengaturan & Reset Data", "✓", "—", "—"],
        ],
    )
    add_para(
        doc,
        "Jika Anda membuka halaman yang tidak diizinkan, sistem mengembalikan Anda ke Dashboard "
        "dan menampilkan pesan bahwa akses ditolak.",
    )

    # 6
    add_heading(doc, "6. Tampilan Dashboard (Setelah Login)", 1)
    add_para(doc, "Setelah login, tampilan terdiri dari:", bold=False)
    add_bullet(doc, "Sidebar kiri: menu navigasi (hanya menu yang boleh diakses peran Anda)")
    add_bullet(doc, "Bagian atas: judul halaman, notifikasi (lonceng), bahasa, tema, dan profil pengguna")
    add_bullet(doc, "Area tengah: isi halaman (tabel, formulir, grafik)")

    add_heading(doc, "6.1 Dashboard Utama", 2)
    add_para(doc, "Halaman pertama menampilkan:")
    add_bullet(doc, "Kartu ringkasan: jumlah pasien, janji hari ini, pendapatan, stok rendah (sesuai peran)")
    add_bullet(doc, "Diagram alur pasien: Registrasi → Janji Temu → Perawatan → Billing → Selesai")
    add_bullet(doc, "Grafik: pertumbuhan pasien, pendapatan bulanan, status janji, jenis perawatan")
    add_bullet(doc, "Tabel janji temu terbaru")

    add_heading(doc, "6.2 Keluar (Logout)", 2)
    add_bullet(doc, "Klik foto/nama di kanan atas → pilih “Keluar”")
    add_bullet(doc, "Anda kembali ke halaman login")

    # 7
    add_heading(doc, "7. Alur Bisnis Klinik (Cara Kerja Ideal)", 1)
    add_para(
        doc,
        "Sistem ini mengikuti alur standar klinik gigi. Gunakan urutan berikut agar data rapi:",
    )
    add_para(doc, "Langkah 1 — Registrasi Pasien", bold=True)
    add_bullet(doc, "Menu: Pasien → Tambah → isi nama, telepon, dll. → Simpan")
    add_para(doc, "Langkah 2 — Janji Temu", bold=True)
    add_bullet(doc, "Menu: Janji Temu → Tambah → pilih pasien, dokter, perawatan, tanggal & jam → Simpan")
    add_bullet(doc, "Status janji: Menunggu → Diproses → Selesai (atau Dibatalkan jika batal)")
    add_para(doc, "Langkah 3 — Perawatan", bold=True)
    add_bullet(doc, "Saat pasien datang, perbarui status janji menjadi Diproses lalu Selesai setelah tindakan selesai")
    add_bullet(doc, "Master layanan di menu Perawatan berisi daftar tindakan dan harga")
    add_para(doc, "Langkah 4 — Billing", bold=True)
    add_bullet(doc, "Menu: Billing → buat invoice untuk pasien → catat total dan pembayaran")
    add_bullet(doc, "Status pembayaran: Belum Bayar, Sebagian, atau Lunas")
    add_para(doc, "Langkah 5 — Selesai & Tindak Lanjut", bold=True)
    add_bullet(doc, "Data tersimpan di sistem; janji berikutnya bisa dijadwalkan kembali dari menu Janji Temu")

    # 8
    add_heading(doc, "8. Panduan per Menu", 1)

    add_heading(doc, "8.1 Pasien", 2)
    add_bullet(doc, "Melihat daftar semua pasien terdaftar")
    add_bullet(doc, "Tambah: tombol hijau/biru “Tambah” → isi formulir → Simpan")
    add_bullet(doc, "Edit: ikon pensil pada baris pasien")
    add_bullet(doc, "Hapus: ikon tempat sampah → konfirmasi di jendela popup (bukan alert browser)")

    add_heading(doc, "8.2 Janji Temu", 2)
    add_bullet(doc, "Mengatur jadwal kunjungan pasien ke dokter")
    add_bullet(doc, "Pilih pasien, dokter, dan jenis perawatan dari daftar yang sudah ada")
    add_bullet(doc, "Perbarui status saat pasien datang atau selesai diperiksa")

    add_heading(doc, "8.3 Perawatan", 2)
    add_bullet(doc, "Daftar layanan klinik (misalnya scaling, tambal gigi, bleaching)")
    add_bullet(doc, "Setiap layanan punya kategori, harga, dan durasi perkiraan")
    add_bullet(doc, "Admin/Super Admin mengelola master ini; dokter biasanya hanya memakai saat membuat janji")

    add_heading(doc, "8.4 Dokter & Staff", 2)
    add_bullet(doc, "Hanya untuk Admin dan Super Admin")
    add_bullet(doc, "Mencatat nama dokter, spesialisasi, jadwal praktik, dan status (Aktif/Cuti)")

    add_heading(doc, "8.5 Inventori", 2)
    add_bullet(doc, "Mencatat stok bahan habis pakai (sarung tangan, masker, dll.)")
    add_bullet(doc, "Sistem memberi peringatan jika jumlah di bawah stok minimum")
    add_bullet(doc, "Notifikasi stok rendah muncul di ikon lonceng (untuk Admin/Super Admin)")

    add_heading(doc, "8.6 Billing", 2)
    add_bullet(doc, "Membuat tagihan (invoice) untuk pasien")
    add_bullet(doc, "Mencatat total tagihan dan jumlah yang sudah dibayar")
    add_bullet(doc, "Memantau tagihan yang belum lunas dari dashboard dan notifikasi")

    add_heading(doc, "8.7 Laporan", 2)
    add_bullet(doc, "Ringkasan: total pasien, pendapatan, piutang")
    add_bullet(doc, "Grafik pendapatan per jenis layanan dan jumlah pasien per gender")
    add_bullet(doc, "Daftar inventori dengan penanda stok rendah")

    add_heading(doc, "8.8 Pengaturan (Super Admin saja)", 2)
    add_bullet(doc, "Mengganti tema tampilan (terang / gelap / ikuti sistem)")
    add_bullet(doc, "Melihat profil akun yang sedang login")
    add_bullet(doc, "Reset Data ke Default: menghapus semua perubahan dan mengembalikan data contoh — hati-hati, tidak bisa dibatalkan")

    # 9
    add_heading(doc, "9. Notifikasi", 1)
    add_para(doc, "Klik ikon lonceng di kanan atas untuk melihat:")
    add_bullet(doc, "Janji temu hari ini")
    add_bullet(doc, "Stok bahan rendah (Admin/Super Admin)")
    add_bullet(doc, "Tagihan belum dibayar (Admin/Super Admin)")
    add_bullet(doc, "Janji temu mendatang")
    add_bullet(doc, "Klik “Tandai dibaca” untuk menandai semua sudah dibaca")
    add_bullet(doc, "Klik salah satu notifikasi untuk membuka halaman terkait")

    # 10
    add_heading(doc, "10. Bahasa & Tema", 1)
    add_bullet(doc, "Bahasa: ikon globe → pilih Bahasa Indonesia atau English (berlaku di seluruh halaman)")
    add_bullet(doc, "Tema: ikon matahari/bulan/ monitor — klik untuk berganti mode Terang, Gelap, atau Sistem")

    # 11
    add_heading(doc, "11. Konfirmasi Hapus & Reset", 1)
    add_para(
        doc,
        "Saat menghapus data pasien, janji, invoice, dll., akan muncul jendela konfirmasi "
        "dengan tombol Batal dan Ya, Hapus. Ini untuk mencegah penghapusan tidak sengaja.",
    )
    add_para(
        doc,
        "Reset data di Pengaturan (Super Admin) juga meminta konfirmasi terlebih dahulu.",
    )

    # 12
    add_heading(doc, "12. Tips Penggunaan Sehari-hari", 1)
    add_bullet(doc, "Daftarkan pasien baru sebelum membuat janji temu pertama kali")
    add_bullet(doc, "Pastikan master Perawatan dan Dokter sudah lengkap agar pemilihan di janji temu mudah")
    add_bullet(doc, "Perbarui status janji setiap hari agar dashboard mencerminkan kondisi terkini")
    add_bullet(doc, "Admin mencatat pembayaran di Billing setelah pasien membayar")
    add_bullet(doc, "Cek notifikasi di pagi hari untuk janji dan stok rendah")
    add_bullet(doc, "Gunakan akun Dokter hanya untuk keperluan klinis; urusan keuangan gunakan akun Admin")

    # 13
    add_heading(doc, "13. Pertanyaan Umum", 1)
    add_para(doc, "Apakah data hilang jika browser ditutup?", bold=True)
    add_para(doc, "Tidak, selama data browser tidak dibersihkan. Data tersimpan otomatis di perangkat yang sama.")

    add_para(doc, "Mengapa saya tidak melihat menu Billing?", bold=True)
    add_para(doc, "Kemungkinan Anda login sebagai Dokter. Gunakan akun Admin atau Super Admin untuk akses keuangan.")

    add_para(doc, "Bagaimana mengganti bahasa?", bold=True)
    add_para(doc, "Klik ikon globe di kanan atas navbar atau header dashboard, lalu pilih bahasa.")

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("— Akhir Panduan —\nArcade Dental ERP · Dokumen Panduan Pengguna")
    fr.font.size = Pt(10)
    fr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build()
